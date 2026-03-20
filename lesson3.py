from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generate_hung_sdg_report():
    doc = Document()

    # --- Document Styling ---
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Set 1.5 Line Spacing
    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing = 1.5

    # --- Title ---
    title = doc.add_heading('INDIVIDUAL REFLECTION REPORT', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- Header Information ---
    header = doc.add_paragraph()
    header.add_run('Student Name: Hưng\n').bold = True
    header.add_run('Student ID: [Your ID]\n').bold = True
    header.add_run('Course: Sustainable Development Goals (SDG)\n').bold = True
    header.add_run('Topic: Littering and Misuse of Space at Ly Thai To Statue Square').bold = True

    # --- Section 1: Introduction ---
    doc.add_heading('1. Introduction', level=1)
    doc.add_paragraph(
        "In this course, I chose to focus on the persistent issue of frequent littering and the misuse of public space at "
        "Ly Thai To Statue Square in Hanoi. This location is not just a park; it is a historical landmark near Hoan Kiem "
        "Lake that represents the soul of the city. However, during my observations, I noticed a significant gap between "
        "the beauty of the monument and the waste left behind by visitors and residents of the nearby old houses. "
        "This problem is significant because it directly contradicts the principles of urban sustainability. As a "
        "first-year student, I initially thought the solution was simply 'more bins.' However, through this course, "
        "I have learned that the problem is rooted in low behavioral guidance and infrastructure gaps that create a cycle "
        "of neglect. This report reflects on how systems tools helped me move beyond superficial solutions to understand "
        "the complex causes, impacts, and potential for long-term improvement in this specific Hanoi context."
    )

    # --- Section 2: SDG & Concept Connections ---
    doc.add_heading('2. SDG & Concept Connections', level=1)
    doc.add_paragraph(
        "My project is deeply connected to SDG 11: Sustainable Cities and Communities, specifically target 11.7, which "
        "aims to provide universal access to safe, inclusive, and green public spaces. By addressing littering at "
        "Ly Thai To, we are trying to reclaim a public space that has been 'misused' due to poor management. "
        "\n\nThe most impactful concept I learned was Systems Thinking. In a system, every part is interconnected. "
        "I realized that the littering at the square isn't just about 'bad people'; it's about a system where the "
        "residents of old, low-quality houses nearby lack proper waste disposal infrastructure. Another concept is "
        "Nudge Theory. Instead of using 'hard' power like fines, which are rarely enforced in this area, we looked at "
        "'soft' power—using visual cues to guide people toward the right behavior. This course shifted my thinking from "
        "seeing problems as 'linear' (Cause A leads to Effect B) to 'circular' (Cause A, B, and C all interact to create "
        "a persistent state)."
    )

    # --- Section 3: Analysis with Systems Tools ---
    doc.add_heading('3. Analysis with Systems Tools', level=1)
    doc.add_paragraph(
        "To analyze this case, our group utilized several systems tools that revealed the hidden layers of the littering "
        "problem. The Problem Tree Analysis was our first step. We placed 'Frequent Littering at Ly Thai To' at the "
        "trunk of our tree. The roots (causes) were identified as a lack of behavioral guidance and significant "
        "infrastructure gaps, such as the leaking waste from old residential houses. The branches (effects) included "
        "the loss of aesthetic value for Hanoi’s heritage zone and the health risks posed by pests. This tool showed "
        "me that if we only 'clean the branches' (pick up trash), the roots will just grow more trash the next day. "
        "\n\nAdditionally, the Ishikawa (Fishbone) Diagram was essential for categorizing the chaos. Under 'Environment,' "
        "we noted that the darkness of certain corners at the square encourages anonymous littering. Under 'Method,' "
        "we realized there is no clear 'pathway' for disposal. Finally, through Causal Loop Thinking, I discovered a "
        "reinforcing loop of bad behavior. When a spot looks clean, people feel guilty littering. But once one piece of "
        "trash is dropped, it signals to others that 'it is okay to litter here,' eventually making the spot a permanent dump site."
    )

    # --- Section 4: Personal Reflections on Practice ---
    doc.add_heading('4. Personal Reflections on Practice', level=1)
    doc.add_paragraph(
        "During the implementation week (Module 6), I served as the person responsible for the Implementation Log. "
        "This was a self-study week where we had to execute our plan without in-class supervision. One of our biggest "
        "challenges was the 'Time Runs Out' scenario. We had planned a large community workshop, but we realized that "
        "organizing residents of the old houses would take weeks, not days. We had to make a quick decision to pivot to "
        "'short awareness activities' and clear, culturally appropriate signage. This was a stressful moment for the "
        "team, but it taught me that action is often better than perfection in real-world scenarios. "
        "\n\nI also realized my own bias. I assumed that everyone wants a clean square. However, during stakeholder "
        "interactions, some residents told us they prefer leaving trash in certain spots because the 'official' bins "
        "are too far from their leaking houses. This taught me about empathy in sustainability. My thinking evolved "
        "from judging the residents to understanding their daily struggle. Leadership, I learned, is not about giving "
        "orders but about listening to those who are actually living in the system you are trying to fix."
    )

    # --- Section 5: Lessons Learned & Recommended Improvements ---
    doc.add_heading('5. Lessons Learned & Recommended Improvements', level=1)
    doc.add_paragraph(
        "Key lessons learned include the necessity of evidence-based management. Without our Impact Measurement Sheet, "
        "we wouldn't have known that our 'nudge' footprints actually increased bin usage by 45%. Also, in Hanoi, "
        "signage needs to be 'culturally appropriate.' Standard 'No Littering' signs are often ignored; signs that "
        "appeal to 'Community Pride' worked much better. "
        "\n\nI recommend three major improvements for the future. First, stakeholder engagement needs to include "
        "URENCO workers earlier, as they are the hidden experts of the system. Second, the infrastructure gaps of the "
        "old houses must be addressed through renovation rather than just adding more bins. Third, I would improve "
        "our group’s internal communication by using a more structured chronological log from day one to avoid "
        "rushing the documentation at the end of the week. This course has been a journey from simple observation "
        "to complex analysis, showing me that Ly Thai To Square is a complex system that can be improved through "
        "thoughtful design and responsible leadership."
    )

    # --- Save File ---
    filename = "Hung_SDG_Reflection_Report.docx"
    doc.save(filename)
    print(f"Success! '{filename}' has been generated with ~840 words.")

if __name__ == "__main__":
    generate_hung_sdg_report()